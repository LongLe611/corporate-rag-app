import streamlit as st
import os
from qdrant_client import QdrantClient
from fastembed import TextEmbedding
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
import tempfile
import re
from typing import List, Dict, Tuple
import uuid
from qdrant_client.models import PointStruct
from datetime import datetime

# Page config
st.set_page_config(
    page_title="Corporate Knowledge RAG System",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem 0;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 1rem 2rem;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        background-color: #d1ecf1;
        border-left: 5px solid #17a2b8;
        border-radius: 5px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'uploaded_files_count' not in st.session_state:
    st.session_state.uploaded_files_count = 0

# Load credentials from Streamlit secrets
@st.cache_resource
def load_config():
    return {
        "qdrant_url": st.secrets["qdrant"]["url"],
        "qdrant_api_key": st.secrets["qdrant"]["api_key"],
        "openrouter_api_key": st.secrets["openrouter"]["api_key"]
    }

config = load_config()

# Initialize clients
@st.cache_resource
def init_clients():
    qdrant = QdrantClient(
        url=config["qdrant_url"],
        api_key=config["qdrant_api_key"],
        timeout=180  # Increase timeout to 180 seconds (3 minutes)
    )
    embedding = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
    llm = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=config["openrouter_api_key"]
    )
    return qdrant, embedding, llm

qdrant_client, embedding_model, llm_client = init_clients()

# Document processing functions
def extract_text_from_pdf(file_bytes) -> Tuple[str, Dict]:
    """Extract text from PDF bytes"""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        metadata = {
            "total_pages": len(doc),
            "title": doc.metadata.get("title", ""),
        }
        
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                text += f"\n--- Page {page_num} ---\n{page_text}"
        
        doc.close()
        return text.strip(), metadata
    except Exception as e:
        st.error(f"Error extracting PDF: {e}")
        return "", {}

def extract_text_from_docx(file_bytes) -> Tuple[str, Dict]:
    """Extract text from DOCX bytes"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        doc = Document(tmp_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_text.append(row_text)
        
        if table_text:
            text += "\n\n--- Tables ---\n" + "\n".join(table_text)
        
        os.unlink(tmp_path)
        
        metadata = {
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
        }
        return text.strip(), metadata
    except Exception as e:
        st.error(f"Error extracting DOCX: {e}")
        return "", {}

def process_uploaded_file(uploaded_file, domain: str, document_type: str) -> Dict:
    """Process uploaded file and return structured data"""
    file_name = uploaded_file.name
    file_ext = os.path.splitext(file_name)[1].lower()
    file_bytes = uploaded_file.read()
    
    # Extract text
    if file_ext == '.pdf':
        text, metadata = extract_text_from_pdf(file_bytes)
    elif file_ext in ['.docx', '.doc']:
        text, metadata = extract_text_from_docx(file_bytes)
    else:
        st.error(f"Unsupported file type: {file_ext}")
        return None
    
    if not text:
        st.error(f"No text extracted from {file_name}")
        return None
    
    # Clean text
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    return {
        "text": text,
        "metadata": {
            "source": file_name,
            "domain": domain,
            "document_type": document_type,
            "file_type": file_ext,
            **metadata
        }
    }

def chunk_document(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    """Split document into overlapping chunks"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        
        if end < text_length:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            
            if break_point > chunk_size * 0.7:
                chunk = text[start:start + break_point + 1]
                end = start + break_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap
    
    return [c for c in chunks if len(c) > 50]

def upload_to_qdrant(doc_data: Dict, batch_size: int = 100) -> int:
    """Upload document chunks to Qdrant with batching for large files"""
    import time
    
    # Chunk document
    chunks = chunk_document(doc_data["text"])
    
    # Generate embeddings
    embeddings = list(embedding_model.embed(chunks))
    
    # Create points
    points = []
    for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        point = PointStruct(
            id=str(uuid.uuid4()),
            vector=embedding.tolist(),
            payload={
                "text": chunk,
                "chunk_index": idx,
                "total_chunks": len(chunks),
                **doc_data["metadata"]
            }
        )
        points.append(point)
    
    # Upload in batches to avoid timeout
    total_uploaded = 0
    
    for i in range(0, len(points), batch_size):
        batch = points[i:i + batch_size]
        
        # Retry logic for each batch
        max_retries = 3
        for retry in range(max_retries):
            try:
                qdrant_client.upsert(
                    collection_name="corporate_documents",
                    points=batch,
                    wait=True
                )
                total_uploaded += len(batch)
                break  # Success, exit retry loop
            except Exception as e:
                if retry < max_retries - 1:
                    time.sleep(2)  # Wait 2 seconds before retry
                    continue
                else:
                    raise Exception(f"Failed to upload batch after {max_retries} retries: {str(e)}")
    
    return total_uploaded

def rag_query(question: str, domain_filter: str, model_name: str, top_k: int = 5):
    """Execute RAG query with corrected Qdrant API"""
    # Generate query embedding
    query_embedding = list(embedding_model.embed([question]))[0]
    
    # Search in Qdrant with corrected API
    if domain_filter and domain_filter != "All Domains":
        results = qdrant_client.query_points(
            collection_name="corporate_documents",
            query=query_embedding.tolist(),
            query_filter={
                "must": [{"key": "domain", "match": {"value": domain_filter}}]
            },
            limit=top_k
        ).points
    else:
        results = qdrant_client.query_points(
            collection_name="corporate_documents",
            query=query_embedding.tolist(),
            limit=top_k
        ).points
    
    if not results:
        return "No relevant documents found.", []
    
    # Build context
    context = "\n\n".join([
        f"[Source: {r.payload.get('source')}]\n{r.payload.get('text')}" 
        for r in results
    ])
    
    # Generate answer
    prompt = f"""Based on the following corporate documents, please answer the question.

Context from documents:
{context}

Question: {question}

Please provide a comprehensive answer based on the information in the documents. If the documents don't contain enough information, please say so. Cite the source documents."""

    response = llm_client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )
    
    answer = response.choices[0].message.content
    sources = [(r.payload.get('source'), r.score) for r in results]
    
    return answer, sources

# App Header
st.markdown('<div class="main-header">📚 Corporate Knowledge RAG System</div>', unsafe_allow_html=True)

# Create tabs
tab1, tab2, tab3 = st.tabs(["💬 Query Documents", "📤 Upload Documents", "📊 Collection Stats"])

# TAB 1: Query Documents
with tab1:
    st.header("Ask Questions About Your Documents")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        question = st.text_input(
            "Enter your question:",
            placeholder="e.g., Chế độ kế toán TT99 năm 2025 quy định những gì?",
            key="question_input"
        )
    
    with col2:
        domain_filter = st.selectbox(
            "Domain Filter:",
            ["All Domains", "Corporate Tax", "Accounting & Finance", "Corporate Law & Regulation"],
            key="domain_filter"
        )
    
    col3, col4, col5 = st.columns(3)
    
    with col3:
        model_choice = st.selectbox(
            "LLM Model:",
            [
                "anthropic/claude-4.5-sonnet",
                "anthropic/claude-4.5-haiku",
                "openai/gpt-5-mini",
                "mistralai/mistral-7b-instruct"
            ],
            key="model_choice"
        )
    
    with col4:
        top_k = st.slider("Number of Sources:", 1, 10, 5, key="top_k")
    
    with col5:
        st.write("")  # Spacing
        query_button = st.button("🔍 Search", type="primary", use_container_width=True)
    
    if query_button and question:
        with st.spinner("Searching and generating answer..."):
            try:
                answer, sources = rag_query(question, domain_filter, model_choice, top_k)
                
                # Display answer
                st.markdown("### 📝 Answer:")
                st.markdown(f'<div class="info-box">{answer}</div>', unsafe_allow_html=True)
                
                # Display sources
                st.markdown("### 📚 Sources Used:")
                for i, (source, score) in enumerate(sources, 1):
                    st.write(f"{i}. **{source}** (Relevance: {score:.4f})")
                
                # Add to history
                st.session_state.chat_history.append({
                    "question": question,
                    "answer": answer,
                    "sources": sources,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
                
            except Exception as e:
                st.error(f"Error: {e}")
    
    # Show chat history
    if st.session_state.chat_history:
        st.markdown("---")
        st.markdown("### 📜 Query History")
        for i, item in enumerate(reversed(st.session_state.chat_history[-5:]), 1):
            with st.expander(f"Q{len(st.session_state.chat_history)-i+1}: {item['question'][:80]}..."):
                st.write(f"**Time:** {item['timestamp']}")
                st.write(f"**Q:** {item['question']}")
                st.write(f"**A:** {item['answer']}")
                st.write("**Sources:**")
                for source, score in item['sources']:
                    st.write(f"  - {source} ({score:.4f})")

# TAB 2: Upload Documents
with tab2:
    st.header("Upload New Documents")
    
    col1, col2 = st.columns(2)
    
    with col1:
        upload_domain = st.selectbox(
            "Select Domain:",
            ["Corporate Tax", "Accounting & Finance", "Corporate Law & Regulation"],
            key="upload_domain"
        )
    
    with col2:
        upload_doc_type = st.selectbox(
            "Document Type:",
            ["policy", "guideline", "regulation", "standard", "memo", "procedure", "general"],
            key="upload_doc_type"
        )
    
    uploaded_files = st.file_uploader(
        "Choose PDF or DOCX files:",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="file_uploader"
    )
    
    if uploaded_files:
        st.info(f"📁 {len(uploaded_files)} file(s) selected")
        
        if st.button("📤 Upload to Qdrant", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            successful = 0
            failed = 0
            total_chunks = 0
            
            for i, uploaded_file in enumerate(uploaded_files):
                # Update progress
                progress_percent = i / len(uploaded_files)
                progress_bar.progress(progress_percent)
                status_text.text(f"Processing {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
                
                try:
                    # Show file size info
                    file_size_kb = len(uploaded_file.getvalue()) / 1024
                    if file_size_kb > 500:  # If file > 500KB, show warning
                        st.info(f"⏳ Large file detected ({file_size_kb:.1f}KB). This may take a moment...")
                    
                    doc_data = process_uploaded_file(uploaded_file, upload_domain, upload_doc_type)
                    
                    if doc_data:
                        chunks_count = upload_to_qdrant(doc_data)
                        total_chunks += chunks_count
                        successful += 1
                        st.success(f"✅ {uploaded_file.name} ({chunks_count} chunks)")
                    else:
                        failed += 1
                        st.error(f"❌ {uploaded_file.name} - Failed to process")
                
                except Exception as e:
                    failed += 1
                    st.error(f"❌ {uploaded_file.name} - Error: {e}")
                
                progress_bar.progress((i + 1) / len(uploaded_files))
            
            status_text.empty()
            progress_bar.empty()
            
            # Summary
            st.markdown(f"""
            <div class="success-box">
                <h4>📊 Upload Summary</h4>
                <ul>
                    <li>✅ Successful: {successful}/{len(uploaded_files)} files</li>
                    <li>❌ Failed: {failed}/{len(uploaded_files)} files</li>
                    <li>📦 Total chunks created: {total_chunks}</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            st.session_state.uploaded_files_count += successful

# TAB 3: Collection Stats
with tab3:
    st.header("Collection Statistics")
    
    if st.button("🔄 Refresh Stats"):
        try:
            collection_info = qdrant_client.get_collection("corporate_documents")
            
            st.metric("Total Vectors", collection_info.points_count)
            st.metric("Vector Dimension", collection_info.config.params.vectors.size)
            
            st.markdown("---")
            st.subheader("📁 Documents by Domain")
            
            for domain in ["Corporate Tax", "Accounting & Finance", "Corporate Law & Regulation"]:
                results = qdrant_client.scroll(
                    collection_name="corporate_documents",
                    scroll_filter={
                        "must": [{"key": "domain", "match": {"value": domain}}]
                    },
                    limit=1000
                )
                
                unique_sources = set([point.payload.get("source") for point in results[0]])
                
                with st.expander(f"📂 {domain} ({len(results[0])} chunks, {len(unique_sources)} documents)"):
                    for source in sorted(unique_sources):
                        st.write(f"• {source}")
        
        except Exception as e:
            st.error(f"Error fetching stats: {e}")